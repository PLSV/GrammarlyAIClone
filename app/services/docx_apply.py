# app/services/docx_apply.py
from typing import List, Dict, Any
from docx import Document
from docx.table import _Cell, Table
import logging

log = logging.getLogger("docx_apply")


def _is_list_paragraph(p) -> bool:
    pPr = getattr(p._p, "pPr", None)
    return bool(pPr is not None and getattr(pPr, "numPr", None) is not None) or ("List" in (p.style.name or ""))


def _iter_paragraphs_in_document(doc: Document):
    # top-level paragraphs
    for p in doc.paragraphs:
        yield p

    # paragraphs inside tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

    # NOTE: paragraphs inside text boxes (w:txbxContent) are NOT exposed by python-docx.
    # Handling those requires raw XML traversal (lxml) and custom mapping. For now we warn.
    try:
        # very rough check if text boxes exist
        xml = doc.element.body.xml
        if "txbxContent" in xml:
            log.warning("DOCX appears to contain text boxes (w:txbxContent). "
                        "python-docx cannot read/update these; content may be skipped.")
    except Exception:
        pass


def collect_segments(docx_path: str) -> List[Dict[str, Any]]:
    """
    Returns a list like:
    [{"index": 0, "kind": "paragraph" | "list_item", "text": "..."}]
    Includes top-level paragraphs and paragraphs inside tables.
    (Text boxes are not supported by python-docx and will be skipped.)
    """
    doc = Document(docx_path)
    segments: List[Dict[str, Any]] = []
    idx = 0

    for p in _iter_paragraphs_in_document(doc):
        text = (p.text or "").strip()
        kind = "list_item" if _is_list_paragraph(p) else "paragraph"
        segments.append({"index": idx, "kind": kind, "text": text})
        idx += 1

    return segments


def apply_rewrites(docx_path: str, rewrites: List[Dict[str, Any]], out_path: str) -> str:
    """
    Replaces paragraph text ONLY (preserves bullet/numbering/indent via style & numbering props).
    Applies to top-level paragraphs and paragraphs inside tables.
    (Text boxes are not supported.)
    """
    doc = Document(docx_path)
    by_idx = {int(r["index"]): str(r["text"]) for r in rewrites if "index" in r and "text" in r}

    idx = 0
    for p in _iter_paragraphs_in_document(doc):
        if idx in by_idx:
            p.text = by_idx[idx]
        idx += 1

    doc.save(out_path)
    return out_path
